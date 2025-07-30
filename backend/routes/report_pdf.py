# backend/routes/report_pdf.py
from fastapi import APIRouter, Request, Response
from fastapi.responses import StreamingResponse
from docx import Document
from io import BytesIO
from datetime import datetime

router = APIRouter()

@router.post("/report/pdf")
async def generate_pdf(request: Request):
    data = await request.json()

    doc = Document()
    doc.add_heading("PreContract Site Assessment Report", 0)
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    doc.add_heading("Project Info", level=1)
    doc.add_paragraph(f"Project Name: {data.get('projectName')}")
    doc.add_paragraph(f"Client: {data.get('firstName')} {data.get('lastName')}")
    doc.add_paragraph(f"Address: {data.get('street')}, {data.get('suburb')} {data.get('state')} {data.get('postcode')}")

    doc.add_heading("Geospatial Metadata", level=1)
    doc.add_paragraph(f"Council: {data.get('council')}")
    doc.add_paragraph(f"Elevation: {data.get('elevation')} m")
    doc.add_paragraph(f"Distance to Coast: {data.get('distanceToCoast')} km")
    doc.add_paragraph(f"Wind Zone: {data.get('windZone')}")
    doc.add_paragraph(f"BAL Rating: {data.get('balRating')}")
    doc.add_paragraph(f"Benchmarks: {data.get('benchmark1')}, {data.get('benchmark2')}")

    doc.add_heading("Assessment", level=1)
    doc.add_paragraph(f"Footing Recommendation: {data.get('footingRecommendation')}")
    doc.add_paragraph(f"Risk Summary: {data.get('riskSummary')}")
    doc.add_paragraph(f"Services Requested: {', '.join(data.get('services', []))}")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    headers = {
        'Content-Disposition': 'inline; filename="precontract-site-report.docx"'
    }
    return StreamingResponse(buffer, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', headers=headers)
